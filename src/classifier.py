from werkzeug.datastructures import FileStorage
import PyPDF2
from docx import Document
import pandas as pd
from PIL import Image
import pytesseract
import pillow_heif
import os
from transformers import pipeline
from functools import partial
import asyncio
from concurrent.futures import ThreadPoolExecutor

classifier = pipeline("zero-shot-classification", model="facebook/bart-large-mnli", device=-1)
executor = ThreadPoolExecutor()

def extract_text_from_pdf(file):
    text = ""
    pdf_reader = PyPDF2.PdfReader(file)
    for page in pdf_reader.pages:
        text += page.extract_text() or ""
    return text

def extract_text_from_csv(file):
    text = ""
    df = pd.read_csv(file)
    for row in df.values:
        text += "\t".join(map(str, row)) + "\n"
    return text

def extract_text_from_docx(file: FileStorage):
    doc = Document(file)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + "\t"
            text += "\n"
    return text

def extract_text_from_excel(file: FileStorage):
    try:
        excel_file = pd.ExcelFile(file)
        text = ""
        for sheet_name in excel_file.sheet_names:
            text += f"Sheet: {sheet_name}\n"
            df = excel_file.parse(sheet_name)
            for _, row in df.iterrows():
                row_text = "\t".join([str(cell) for cell in row])
                text += row_text + "\n"
        return text
    except Exception as e:
        return f"Error processing Excel file: {str(e)}"
    
def extract_text_from_image(filename: str):
    try:
        if filename.lower().endswith('.heic'):
            heif_file = pillow_heif.read_heif(filename)
            image = Image.frombytes(
                heif_file.mode,
                heif_file.size,
                heif_file.data,
                "raw",
            )
            image.save("temp.png", format="png")
            image = Image.open("temp.png")
            text = pytesseract.image_to_string(image)
            os.remove("temp.png")
        else:
            image = Image.open(filename)
            text = pytesseract.image_to_string(image)
        return text
    except Exception as e:
        return f"Error processing image: {str(e)}"

def run_extraction(extraction_func, file_stream):
    """Run the extraction function in a separate thread"""
    return extraction_func(file_stream)

def run_classification(text, labels):
    """Run the classification in a separate thread"""
    return classifier(text, labels)

async def classify_file(file: FileStorage):
    if isinstance(file, FileStorage):
        file_extension = file.filename.lower()
        file_stream = file.stream
    elif isinstance(file, str) and os.path.isfile(file):
        file_extension = file.lower()
        file_stream = open(file, 'rb')
    else:
        return "Invalid file input"

    extraction_functions = {
        '.pdf': extract_text_from_pdf,
        '.docx': extract_text_from_docx,
        '.png': extract_text_from_image,
        '.jpg': extract_text_from_image,
        '.jpeg': extract_text_from_image,
        '.heic': extract_text_from_image,
        '.xlsx': extract_text_from_excel,
        '.csv': extract_text_from_csv,
    }

    extraction_func = next(
        (func for ext, func in extraction_functions.items() if file_extension.endswith(ext)), None
    )

    if not extraction_func:
        return "File type not recognised"
    
    try:
        # Run the extraction in a thread pool
        loop = asyncio.get_event_loop()
        text = await loop.run_in_executor(executor, run_extraction, extraction_func, file_stream)
        
        if isinstance(file_stream, FileStorage):
            file_stream.close()
            
        if not text.strip():
            return "No text extracted from the file"

        candidate_labels = ["invoice", "bank statement", "drivers license", "passport", "credit note", "cv", "unknown file"]
        
        result = await loop.run_in_executor(executor, run_classification, text, candidate_labels)
        
        return result['labels'][0]
        
    except Exception as e:
        return f"Error processing file: {str(e)}"
